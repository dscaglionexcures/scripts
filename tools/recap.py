#!/usr/bin/env python3
"""
Evaluate a checklist for a subject and generate a polished PDF report.

Features:
- Calls POST /api/v1/patient-registry/checklist/{checklistId}/evaluate
- Resolves checklist display name from GET /api/v1/patient-registry/checklist
- Renders a sectioned PDF that adapts to arbitrary checklist item result schemas
- Optionally saves the raw evaluation JSON alongside the PDF
"""

from __future__ import annotations

import argparse
import json
import os
import re
import sys
from dataclasses import dataclass
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import Any, Dict, List, Optional, Sequence
from xml.sax.saxutils import escape
from zoneinfo import ZoneInfo

import requests
from pypdf import PdfReader, PdfWriter
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from xcures_toolkit.api_common import require_json_list, require_json_object
from xcures_toolkit.auth_common import get_xcures_bearer_token, load_env_file
from xcures_toolkit.xcures_client import XcuresApiClient


DEFAULT_BASE_URL = "https://partner.xcures.com"
REPO_ROOT = Path(__file__).resolve().parent.parent
DEFAULT_OUTPUT_DIR = REPO_ROOT / "downloads"
EVALUATE_ENDPOINT_TEMPLATE = "/api/v1/patient-registry/checklist/{checklist_id}/evaluate"
VITALS_ENDPOINT = "/api/v1/patient-registry/clinical-concepts/vital"
LABS_ENDPOINT = "/api/v1/patient-registry/clinical-concepts/lab"
TEMPLATE_ORG_NAME = "Vitality Consultants, LLC"
TEMPLATE_ORG_ADDRESS = "3540 Toringdon Way Ste 200, Charlotte, NC 28277-4650"
TEMPLATE_TITLE_LINE_1 = "Comprehensive Clinical Consultation"
TEMPLATE_TITLE_LINE_2 = "Summary Note"


@dataclass(frozen=True)
class DetailRow:
    label: str
    value: str


@dataclass(frozen=True)
class ItemViewModel:
    title: str
    sort_order: float
    meets_criteria: bool
    details: List[DetailRow]
    result_raw: Optional[Dict[str, Any]]
    evidence: str
    source_lines: List[str]
    document_count: int
    result_is_structured: bool


@dataclass(frozen=True)
class ReportMeta:
    subject_id: str
    subject_full_name: str
    subject_external_id: str
    subject_dob: str
    checklist_id: str
    checklist_name: str
    eligibility_satisfied: bool
    generated_at: datetime
    total_items: int
    passed_items: int
    failed_items: int
    unique_document_count: int
    endpoint_path: str


@dataclass(frozen=True)
class RenderConfig:
    page_size: tuple[float, float]
    margin_left: float
    margin_right: float
    margin_top: float
    margin_bottom: float
    header_height: float
    footer_height: float
    accent_color: colors.Color
    header_bg_color: colors.Color
    pass_color: colors.Color
    fail_color: colors.Color
    border_color: colors.Color
    muted_text_color: colors.Color
    card_bg_color: colors.Color


def build_render_config() -> RenderConfig:
    return RenderConfig(
        page_size=(8.5 * inch, 11.5 * inch),
        margin_left=0.6 * inch,
        margin_right=0.6 * inch,
        margin_top=2.35 * inch,
        margin_bottom=0.95 * inch,
        header_height=2.1 * inch,
        footer_height=0.55 * inch,
        accent_color=colors.HexColor("#000000"),
        header_bg_color=colors.HexColor("#FFFFFF"),
        pass_color=colors.HexColor("#2E7D32"),
        fail_color=colors.HexColor("#B71C1C"),
        border_color=colors.HexColor("#A6A6A6"),
        muted_text_color=colors.HexColor("#2B2B2B"),
        card_bg_color=colors.HexColor("#FFFFFF"),
    )


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description=(
            "Evaluate a checklist for a subject and export a polished PDF report "
            "(optionally with raw JSON sidecar)."
        )
    )
    parser.add_argument(
        "--subject-id",
        required=False,
        help="Subject UUID to evaluate. If omitted, you will be prompted.",
    )
    parser.add_argument(
        "--checklist-id",
        required=False,
        help="Checklist UUID to evaluate. If omitted, you will pick from a numbered list.",
    )
    parser.add_argument(
        "--project-id",
        default=os.getenv("XCURES_PROJECT_ID", "").strip() or None,
        help="ProjectId header value (defaults to XCURES_PROJECT_ID env var).",
    )
    parser.add_argument(
        "--base-url",
        default=(
            os.getenv("XCURES_BASE_URL", "").strip()
            or os.getenv("BASE_URL", "").strip()
            or DEFAULT_BASE_URL
        ),
        help="API base URL.",
    )
    parser.add_argument(
        "--bearer",
        default=None,
        help="Optional bearer token override. If absent, uses client credentials flow.",
    )
    parser.add_argument(
        "--regenerate",
        action="store_true",
        help="Force checklist re-evaluation even if cached results exist.",
    )
    parser.add_argument(
        "--save-json",
        dest="save_json",
        action="store_true",
        default=True,
        help="Save raw evaluation JSON next to the PDF output (default: enabled).",
    )
    parser.add_argument(
        "--no-save-json",
        dest="save_json",
        action="store_false",
        help="Disable saving raw evaluation JSON sidecar for this run.",
    )
    parser.add_argument(
        "--save-docx",
        dest="save_docx",
        action="store_true",
        default=True,
        help="Save Word (.docx) output next to the PDF output (default: enabled).",
    )
    parser.add_argument(
        "--no-save-docx",
        dest="save_docx",
        action="store_false",
        help="Disable Word (.docx) export for this run.",
    )
    parser.add_argument(
        "--output-dir",
        default=str(DEFAULT_OUTPUT_DIR),
        help=(
            "Ignored. Exports always go to "
            f"{DEFAULT_OUTPUT_DIR}."
        ),
    )
    parser.add_argument(
        "--timeout",
        type=int,
        default=60,
        help="Request timeout in seconds (default: 60).",
    )
    args = parser.parse_args()
    if not args.project_id:
        parser.error("--project-id is required (or set XCURES_PROJECT_ID in the environment).")
    return args


def normalize_label(key: str) -> str:
    cleaned = key.replace("_", " ").replace("-", " ").strip()
    if not cleaned:
        return "Field"
    return " ".join(part.capitalize() for part in cleaned.split())


def _is_scalar(value: Any) -> bool:
    return not isinstance(value, (dict, list))


def _indent_multiline(text: str, prefix: str = "  ") -> str:
    lines = text.splitlines()
    if not lines:
        return text
    return "\n".join(f"{prefix}{line}" for line in lines)


def _scalar_to_text(value: Any) -> str:
    if value is None:
        return "N/A"
    if isinstance(value, bool):
        return "true" if value else "false"
    return str(value)


def _compact_dict_values(value: Dict[str, Any]) -> str:
    if not value:
        return ""
    if not all(_is_scalar(v) for v in value.values()):
        return ""

    priority = [
        "medication",
        "dose",
        "formulation",
        "patient_instructions",
        "reason",
        "allergen",
        "reaction",
        "recorded_date",
        "surgical_procedure",
        "procedure_date",
        "anatomic_location",
        "notes",
        "condition",
        "relationship",
        "diagnostic_procedure",
        "findings",
        "impression",
        "visit_date",
        "vision_visit",
        "dental_visit",
        "immunization",
        "immunization_date",
        "justification",
    ]

    normalized_to_raw: Dict[str, str] = {}
    for raw_key in value.keys():
        normalized_to_raw[re.sub(r"[^a-z0-9]", "", str(raw_key).lower())] = str(raw_key)

    ordered_keys: List[str] = []
    for candidate in priority:
        match_key = normalized_to_raw.get(re.sub(r"[^a-z0-9]", "", candidate.lower()))
        if match_key and match_key not in ordered_keys:
            ordered_keys.append(match_key)
    for raw_key in value.keys():
        raw_key_str = str(raw_key)
        if raw_key_str not in ordered_keys:
            ordered_keys.append(raw_key_str)

    parts: List[str] = []
    for raw_key in ordered_keys:
        raw_val = value.get(raw_key)
        text = _scalar_to_text(raw_val).strip()
        if not text or text == "N/A":
            continue
        parts.append(text)

    return ", ".join(parts)


def _humanize_value(value: Any, *, depth: int = 0) -> str:
    if depth > 6:
        return "..."

    if _is_scalar(value):
        return _scalar_to_text(value)

    if isinstance(value, dict):
        if not value:
            return "N/A"

        compact_line = _compact_dict_values(value)
        if compact_line:
            return compact_line

        lines: List[str] = []
        for key, nested_value in value.items():
            label = normalize_label(str(key))
            nested_text = _humanize_value(nested_value, depth=depth + 1)
            if "\n" in nested_text:
                lines.append(f"{label}:")
                lines.append(_indent_multiline(nested_text, "  "))
            else:
                lines.append(f"{label}: {nested_text}")
        return "\n".join(lines)

    # value is a list at this point
    if not value:
        return "N/A"

    if all(_is_scalar(item) for item in value):
        return ", ".join(_scalar_to_text(item) for item in value)

    flattened_lines: List[str] = []
    for item in value:
        nested_text = _humanize_value(item, depth=depth + 1)
        nested_parts = [part.strip() for part in nested_text.splitlines() if part.strip()]
        flattened_lines.extend(nested_parts or [nested_text.strip()])
    return "\n".join(line for line in flattened_lines if line)


def value_to_text(value: Any) -> str:
    return _humanize_value(value)


def compact_id(identifier: str, max_len: int = 8) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9]+", "", identifier or "")
    if not cleaned:
        return "id"
    return cleaned[:max_len]


def html_paragraph_text(raw: str) -> str:
    normalized = (raw or "")
    for glyph in ("•", "◦", "▪", "‣", "●", "∙", "·", "\uf0b7", "\ufffd"):
        normalized = normalized.replace(glyph, "-")
    safe = escape(normalized)
    return safe.replace("\n", "<br/>")


def _bold_colon_label_in_line(line: str) -> str:
    # For human-facing key/value lines, bold only short label-like prefixes.
    # This avoids over-formatting prose that contains references like
    # "(demographic:5, social_history:6)".
    #
    # Examples matched:
    # - "Patient Instructions: take with food" -> "<b>Patient Instructions:</b> ..."
    # - "- Medication: aspirin 81 mg" -> "- <b>Medication:</b> ..."
    # - "Alcohol Use Status:" -> "<b>Alcohol Use Status:</b>"
    safe = escape(line)
    match = re.match(
        r"^(\s*(?:-\s+)?)"
        r"((?:[A-Za-z][A-Za-z0-9'()/.\-]*"
        r"(?:\s+[A-Za-z][A-Za-z0-9'()/.\-]*){0,5}):)"
        r"(\s*.*)$",
        safe,
    )
    if not match:
        return safe

    prefix, label, suffix = match.groups()
    return f"{prefix}<b>{label}</b>{suffix}"


def html_paragraph_text_with_colon_word_bold(raw: str) -> str:
    normalized = (raw or "")
    for glyph in ("•", "◦", "▪", "‣", "●", "∙", "·", "\uf0b7", "\ufffd"):
        normalized = normalized.replace(glyph, "-")
    lines = normalized.splitlines()
    if not lines:
        return ""
    return "<br/>".join(_bold_colon_label_in_line(line) for line in lines)


def styled_paragraph(raw: str, style: ParagraphStyle) -> Paragraph:
    return Paragraph(html_paragraph_text(raw), style)


def styled_paragraph_with_colon_word_bold(raw: str, style: ParagraphStyle) -> Paragraph:
    return Paragraph(html_paragraph_text_with_colon_word_bold(raw), style)


def append_text_with_colon_headings(
    block: List[Any],
    text: str,
    *,
    body_style: ParagraphStyle,
) -> None:
    lines = (text or "").splitlines()
    if not lines:
        return

    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()
        if not stripped:
            block.append(Spacer(1, 0.02 * inch))
            continue
        if ":" in stripped:
            block.append(styled_paragraph(line, body_style))
            if stripped.endswith(":"):
                block.append(Spacer(1, 0.02 * inch))
            continue
        block.append(styled_paragraph(line, body_style))


def extract_evidence(result: Any) -> str:
    if not isinstance(result, dict):
        return "No structured result provided."
    justification = result.get("justification")
    if isinstance(justification, str) and justification.strip():
        return justification.strip()
    evidence = result.get("evidence")
    if evidence is None:
        return "No evidence text provided."
    return value_to_text(evidence)


def build_source_lines(item: Dict[str, Any]) -> tuple[List[str], int]:
    lines: List[str] = []
    document_ids = item.get("documentIds")
    if not isinstance(document_ids, list):
        document_ids = []

    section_doc_ids: set[str] = set()
    sections = item.get("documentSections")
    if isinstance(sections, list) and sections:
        for section in sections:
            if not isinstance(section, dict):
                continue
            doc_id = str(section.get("documentId") or "").strip()
            if doc_id:
                section_doc_ids.add(doc_id)

            doc_name = str(section.get("documentName") or "Document").strip()
            doc_date = str(section.get("documentDate") or "date unknown").strip()
            section_type = str(section.get("sectionType") or "").strip()
            section_ref = str(section.get("section") or "").strip()
            section_title = str(section.get("sectionTitle") or "").strip()

            section_bits = [f"type={section_type or 'unknown'}"]
            if section_ref:
                section_bits.append(f"ref={section_ref}")
            if section_title and section_title.lower() != "none":
                section_bits.append(f"title={section_title}")

            lines.append(f"{doc_name} ({doc_date}) - {', '.join(section_bits)}")
    else:
        for doc_id in document_ids[:12]:
            lines.append(f"documentId={doc_id}")
        if len(document_ids) > 12:
            lines.append(f"... and {len(document_ids) - 12} more document IDs")

    unique_document_count = len(set(str(x) for x in document_ids if x) | section_doc_ids)
    if not lines:
        lines = ["No source references provided."]
    return lines, unique_document_count


def normalize_item(item: Dict[str, Any]) -> ItemViewModel:
    checklist_item = item.get("checklistItem")
    checklist_item_obj = checklist_item if isinstance(checklist_item, dict) else {}
    title = (
        str(checklist_item_obj.get("name") or "").strip()
        or str(checklist_item_obj.get("libraryItemId") or "").strip()
        or str(checklist_item_obj.get("id") or "").strip()
        or "Untitled checklist item"
    )

    raw_sort_order = checklist_item_obj.get("sortOrder")
    try:
        sort_order = float(raw_sort_order)
    except Exception:
        sort_order = 10_000.0

    meets_criteria = bool(item.get("meetsCriteria"))

    result = item.get("result")
    details: List[DetailRow] = []
    if isinstance(result, dict):
        for key, value in result.items():
            details.append(DetailRow(label=normalize_label(str(key)), value=value_to_text(value)))

    source_lines, doc_count = build_source_lines(item)
    return ItemViewModel(
        title=title,
        sort_order=sort_order,
        meets_criteria=meets_criteria,
        details=details,
        result_raw=result if isinstance(result, dict) else None,
        evidence=extract_evidence(result),
        source_lines=source_lines,
        document_count=doc_count,
        result_is_structured=isinstance(result, dict),
    )


def clean_section_title(title: str, checklist_name: str) -> str:
    raw_title = (title or "").strip()
    if not raw_title:
        return raw_title

    checklist = (checklist_name or "").strip()
    normalized = raw_title
    if checklist:
        prefix_pattern = re.compile(rf"^\s*{re.escape(checklist)}\s*-\s*", re.IGNORECASE)
        cleaned = prefix_pattern.sub("", raw_title, count=1).strip()
        if cleaned:
            normalized = cleaned

    # Fallback for older checklist labels that still include "RECAP - ...".
    if normalized == raw_title:
        recap_cleaned = re.sub(r"^\s*RECAP\s*-\s*", "", raw_title, flags=re.IGNORECASE).strip()
        normalized = recap_cleaned if recap_cleaned else raw_title

    aliases = {
        "conditions and comorbidities": "Patient Medical History",
        "social determinants of health": "SDOH/Mental Health",
    }
    aliased = aliases.get(normalized.lower())
    if aliased:
        return aliased

    return normalized


def build_styles(config: RenderConfig) -> Dict[str, ParagraphStyle]:
    sample = getSampleStyleSheet()
    return {
        "section_label": ParagraphStyle(
            "SectionLabel",
            parent=sample["Heading2"],
            fontName="Helvetica",
            fontSize=10.9,
            leading=12.8,
            textColor=colors.black,
            spaceBefore=0,
            spaceAfter=0,
        ),
        "h1": ParagraphStyle(
            "H1Custom",
            parent=sample["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=11.2,
            leading=13.2,
            textColor=colors.black,
            spaceBefore=5,
            spaceAfter=2,
        ),
        "h2": ParagraphStyle(
            "H2Custom",
            parent=sample["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=10.2,
            leading=12,
            textColor=colors.black,
            spaceBefore=4,
            spaceAfter=2,
        ),
        "h3": ParagraphStyle(
            "H3Custom",
            parent=sample["Heading3"],
            fontName="Helvetica-Bold",
            fontSize=9.8,
            leading=11.4,
            textColor=colors.black,
            spaceBefore=2,
            spaceAfter=1.5,
        ),
        "section": ParagraphStyle(
            "SectionHeading",
            parent=sample["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=10.8,
            leading=12.6,
            textColor=colors.black,
            spaceBefore=5,
            spaceAfter=2,
        ),
        "body": ParagraphStyle(
            "BodyTextCustom",
            parent=sample["BodyText"],
            fontName="Helvetica",
            fontSize=9,
            leading=11.2,
            textColor=colors.black,
            spaceBefore=0.5,
            spaceAfter=0.8,
        ),
        "small": ParagraphStyle(
            "SmallText",
            parent=sample["BodyText"],
            fontName="Helvetica",
            fontSize=8.2,
            leading=10.2,
            textColor=config.muted_text_color,
        ),
        "table_header": ParagraphStyle(
            "TableHeader",
            parent=sample["BodyText"],
            fontName="Helvetica-Bold",
            fontSize=9,
            textColor=colors.black,
            leading=11,
        ),
    }


def draw_page_chrome(
    canvas_obj: Any,
    doc: Any,
    *,
    meta: ReportMeta,
    config: RenderConfig,
) -> None:
    del doc

    def truncate_for_width(text: str, font_name: str, font_size: float, max_width: float) -> str:
        value = text or ""
        if canvas_obj.stringWidth(value, font_name, font_size) <= max_width:
            return value
        ellipsis = "..."
        current = value
        while current and (
            canvas_obj.stringWidth(current + ellipsis, font_name, font_size) > max_width
        ):
            current = current[:-1]
        return (current + ellipsis) if current else ellipsis

    width, height = config.page_size
    generated_display = meta.generated_at.strftime("%Y %b %d %H:%M")

    birthday = "N/A"
    dob_match = re.match(r"^(\d{4}-\d{2}-\d{2})", meta.subject_dob or "")
    if dob_match:
        birthday = dob_match.group(1)

    age_text = "N/A"
    age_match = re.search(r"\((\d+)\s+Year[s]?\s+Old\)", meta.subject_dob or "")
    if age_match:
        age_text = age_match.group(1)

    canvas_obj.saveState()
    canvas_obj.setFillColor(colors.black)

    content_width = width - config.margin_left - config.margin_right
    header_gutter = 0.3 * inch
    column_width = (content_width - header_gutter) / 2
    left_x = config.margin_left
    right_x = config.margin_left + column_width + header_gutter
    top_y = height - 0.42 * inch

    y = top_y
    canvas_obj.setFont("Helvetica-Bold", 11.7)
    canvas_obj.drawString(
        left_x,
        y,
        truncate_for_width(TEMPLATE_ORG_NAME, "Helvetica-Bold", 11.7, column_width),
    )

    y -= 0.16 * inch
    canvas_obj.setFont("Helvetica", 8.7)
    canvas_obj.drawString(
        left_x,
        y,
        truncate_for_width(TEMPLATE_ORG_ADDRESS, "Helvetica", 8.7, column_width),
    )

    canvas_obj.setFont("Helvetica-Bold", 11.3)
    canvas_obj.drawString(
        right_x,
        top_y,
        truncate_for_width(TEMPLATE_TITLE_LINE_1, "Helvetica-Bold", 11.3, column_width),
    )

    canvas_obj.drawString(
        right_x,
        top_y - (0.18 * inch),
        truncate_for_width(TEMPLATE_TITLE_LINE_2, "Helvetica-Bold", 11.3, column_width),
    )

    y = top_y - (0.46 * inch)
    max_text_width = content_width
    canvas_obj.setFont("Helvetica-Bold", 11)
    canvas_obj.drawString(
        config.margin_left,
        y,
        truncate_for_width(meta.subject_full_name or meta.subject_id, "Helvetica-Bold", 11, max_text_width),
    )

    y -= 0.16 * inch
    canvas_obj.setFont("Helvetica", 9)
    canvas_obj.drawString(
        config.margin_left,
        y,
        truncate_for_width(f"MRN : {meta.subject_external_id}", "Helvetica", 9, max_text_width),
    )

    y -= 0.16 * inch
    canvas_obj.drawString(
        config.margin_left,
        y,
        truncate_for_width(
            f"Birthday : {birthday}  Phone : N/A",
            "Helvetica",
            9,
            max_text_width,
        ),
    )

    y -= 0.16 * inch
    canvas_obj.drawString(
        config.margin_left,
        y,
        truncate_for_width(
            f"Visited on: N/A (Age at visit: {age_text} years) "
            f"Electronically signed by: N/A on {generated_display}",
            "Helvetica",
            9,
            max_text_width,
        ),
    )

    canvas_obj.setStrokeColor(config.border_color)
    canvas_obj.line(
        config.margin_left,
        height - config.header_height,
        width - config.margin_right,
        height - config.header_height,
    )
    canvas_obj.restoreState()


def draw_footer(
    canvas_obj: Any,
    *,
    meta: ReportMeta,
    config: RenderConfig,
    page_number: int,
    total_pages: int,
) -> None:
    width, _ = config.page_size
    printed_on = meta.generated_at.strftime("%Y %b %d %H:%M")

    canvas_obj.saveState()
    canvas_obj.setStrokeColor(config.border_color)
    canvas_obj.line(config.margin_left, config.footer_height, width - config.margin_right, config.footer_height)

    canvas_obj.setFillColor(colors.black)
    canvas_obj.setFont("Helvetica", 8.1)
    canvas_obj.drawString(config.margin_left, 0.34 * inch, f"Printed on: {printed_on}")
    canvas_obj.drawRightString(
        width - config.margin_right,
        0.34 * inch,
        f"Page {page_number} of {total_pages}",
    )
    canvas_obj.drawString(config.margin_left, 0.2 * inch, "Note created using Tebra")
    canvas_obj.restoreState()


class NumberedCanvas(canvas.Canvas):
    def __init__(self, *args: Any, meta: ReportMeta, config: RenderConfig, **kwargs: Any) -> None:
        super().__init__(*args, **kwargs)
        self._saved_page_states: List[Dict[str, Any]] = []
        self._meta = meta
        self._config = config

    def showPage(self) -> None:
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self) -> None:
        total_pages = len(self._saved_page_states)
        for page_state in self._saved_page_states:
            self.__dict__.update(page_state)
            draw_footer(
                self,
                meta=self._meta,
                config=self._config,
                page_number=self._pageNumber,
                total_pages=total_pages,
            )
            canvas.Canvas.showPage(self)
        canvas.Canvas.save(self)


def build_item_block(
    *,
    index: int,
    item: ItemViewModel,
    styles: Dict[str, ParagraphStyle],
    config: RenderConfig,
    usable_width: float,
) -> List[Any]:
    del index
    left_col_width = min(max(1.45 * inch, usable_width * 0.24), 1.9 * inch)
    right_col_width = usable_width - left_col_width

    title = Paragraph(f"{html_paragraph_text(item.title)}", styles["section_label"])
    right_column: List[Any] = []

    emitted_text = False
    if item.result_is_structured and item.details:
        for row in item.details:
            value = (row.value or "").strip() or "N/A"
            label_with_colon = f"{row.label}:"
            if row.label.strip().lower() == "items":
                if value != "N/A":
                    append_text_with_colon_headings(right_column, value, body_style=styles["body"])
                emitted_text = True
                continue
            if "\n" in value:
                right_column.append(styled_paragraph_with_colon_word_bold(label_with_colon, styles["body"]))
                right_column.append(Spacer(1, 0.01 * inch))
                append_text_with_colon_headings(right_column, value, body_style=styles["body"])
            else:
                right_column.append(
                    styled_paragraph_with_colon_word_bold(f"{label_with_colon} {value}", styles["body"])
                )
            emitted_text = True
    else:
        evidence = (item.evidence or "").strip()
        if evidence:
            append_text_with_colon_headings(right_column, evidence, body_style=styles["body"])
            emitted_text = True

    if not emitted_text:
        right_column.append(styled_paragraph("No structured result provided.", styles["body"]))

    # Build one table row per flowable so long sections can naturally split across pages.
    # A single-row table with the whole section body in one cell cannot split and triggers
    # LayoutError for larger sections.
    table_rows: List[List[Any]] = []
    for idx, flowable in enumerate(right_column):
        left_cell: Any = title if idx == 0 else Spacer(0, 0)
        table_rows.append([left_cell, flowable])

    section_table = Table(table_rows, colWidths=[left_col_width, right_col_width])
    section_table.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LINEBELOW", (0, -1), (-1, -1), 0.6, config.border_color),
                ("LEFTPADDING", (0, 0), (0, -1), 0),
                ("RIGHTPADDING", (0, 0), (0, -1), 12),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ("LEFTPADDING", (1, 0), (1, -1), 0),
                ("RIGHTPADDING", (1, 0), (1, -1), 0),
            ]
        )
    )

    return [section_table, Spacer(1, 0.11 * inch)]


def write_pdf_report(
    *,
    meta: ReportMeta,
    items: Sequence[ItemViewModel],
    output_pdf: Path,
    config: RenderConfig,
) -> None:
    styles = build_styles(config)
    doc = SimpleDocTemplate(
        str(output_pdf),
        pagesize=config.page_size,
        leftMargin=config.margin_left,
        rightMargin=config.margin_right,
        topMargin=config.margin_top,
        bottomMargin=config.margin_bottom,
        title="Checklist Evaluation Report",
        author="xCures Checklist Export",
    )

    usable_width = config.page_size[0] - config.margin_left - config.margin_right
    story: List[Any] = []
    if not items:
        story.append(Paragraph("No checklist items returned.", styles["body"]))
    else:
        for idx, item in enumerate(items, start=1):
            story.extend(
                build_item_block(
                    index=idx,
                    item=item,
                    styles=styles,
                    config=config,
                    usable_width=usable_width,
                )
            )

    def _on_page(canvas_obj: Any, doc_obj: Any) -> None:
        draw_page_chrome(canvas_obj, doc_obj, meta=meta, config=config)

    def _canvas_maker(*args: Any, **kwargs: Any) -> NumberedCanvas:
        return NumberedCanvas(*args, meta=meta, config=config, **kwargs)

    doc.build(
        story,
        onFirstPage=_on_page,
        onLaterPages=_on_page,
        canvasmaker=_canvas_maker,
    )


def _normalize_word_text(raw: str) -> str:
    normalized = (raw or "")
    for glyph in ("•", "◦", "▪", "‣", "●", "∙", "·", "\uf0b7", "\ufffd"):
        normalized = normalized.replace(glyph, "-")
    return normalized


def _item_word_lines(item: ItemViewModel) -> List[str]:
    lines: List[str] = []
    emitted_text = False

    if item.result_is_structured and item.details:
        for row in item.details:
            value = _normalize_word_text((row.value or "").strip() or "N/A")
            label = (row.label or "Field").strip() or "Field"
            label_with_colon = f"{label}:"

            if label.lower() == "items":
                if value != "N/A":
                    lines.extend(value.splitlines() or [value])
                emitted_text = True
                continue

            if "\n" in value:
                lines.append(label_with_colon)
                lines.extend(value.splitlines())
            else:
                lines.append(f"{label_with_colon} {value}")
            emitted_text = True
    else:
        evidence = _normalize_word_text((item.evidence or "").strip())
        if evidence:
            lines.extend(evidence.splitlines())
            emitted_text = True

    if not emitted_text:
        lines.append("No structured result provided.")
    return lines


def _append_word_line(paragraph: Any, line: str) -> None:
    text = _normalize_word_text(line or "")
    match = re.match(
        r"^(\s*(?:-\s+)?)"
        r"((?:[A-Za-z][A-Za-z0-9'()/.\-]*"
        r"(?:\s+[A-Za-z][A-Za-z0-9'()/.\-]*){0,5}):)"
        r"(\s*.*)$",
        text,
    )
    if not match:
        lowered = text.strip().lower()
        if lowered in {"never smoker", "do not drink"}:
            run = paragraph.add_run(text)
            run.bold = True
            return

        # Mirror the PDF style where list-like entries frequently bold the lead term.
        if "," in text and ":" not in text:
            lead, tail = text.split(",", 1)
            if 1 <= len(lead.strip()) <= 90:
                lead_run = paragraph.add_run(lead)
                lead_run.bold = True
                paragraph.add_run("," + tail)
                return

        paragraph.add_run(text)
        return

    prefix, label, suffix = match.groups()
    if prefix:
        paragraph.add_run(prefix)
    label_run = paragraph.add_run(label)
    label_run.bold = True
    if suffix:
        paragraph.add_run(suffix)


def _set_docx_cell_border(cell: Any, **kwargs: Dict[str, str]) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("left", "top", "right", "bottom", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if edge_data:
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key, value in edge_data.items():
                element.set(qn(f"w:{key}"), str(value))
        elif element is not None:
            tc_borders.remove(element)


def _remove_table_borders(table: Any) -> None:
    for row in table.rows:
        for cell in row.cells:
            _set_docx_cell_border(
                cell,
                left={"val": "nil"},
                right={"val": "nil"},
                top={"val": "nil"},
                bottom={"val": "nil"},
            )


def _set_docx_cell_margins(
    cell: Any,
    *,
    top: int | None = None,
    bottom: int | None = None,
    left: int | None = None,
    right: int | None = None,
) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for tag, value in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        if value is None:
            continue
        element = tc_mar.find(qn(f"w:{tag}"))
        if element is None:
            element = OxmlElement(f"w:{tag}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _add_horizontal_rule(doc: Any) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "A6A6A6")
    p_bdr.append(bottom)


def write_docx_report(
    *,
    meta: ReportMeta,
    items: Sequence[ItemViewModel],
    output_docx: Path,
) -> None:
    try:
        from docx import Document
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is required for Word export. Install it with: "
            ".venv_sdk/bin/pip install python-docx"
        ) from exc

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.7)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Helvetica"
    normal_style.font.size = Pt(10)

    header_table = doc.add_table(rows=2, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    header_table.autofit = False
    header_table.columns[0].width = Inches(3.65)
    header_table.columns[1].width = Inches(3.65)
    _remove_table_borders(header_table)

    left_top = header_table.cell(0, 0)
    left_top.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    p = left_top.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    org_run = p.add_run(TEMPLATE_ORG_NAME)
    org_run.bold = True
    org_run.font.size = Pt(12)

    left_bottom = header_table.cell(1, 0)
    p = left_bottom.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    addr_run = p.add_run(TEMPLATE_ORG_ADDRESS)
    addr_run.font.size = Pt(8)

    right_top = header_table.cell(0, 1)
    p = right_top.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    title_run = p.add_run(f"{TEMPLATE_TITLE_LINE_1}\n{TEMPLATE_TITLE_LINE_2}")
    title_run.bold = True
    title_run.font.size = Pt(12)

    right_bottom = header_table.cell(1, 1)
    right_bottom.text = ""

    _add_horizontal_rule(doc)

    birthday = "N/A"
    dob_match = re.match(r"^(\d{4}-\d{2}-\d{2})", meta.subject_dob or "")
    if dob_match:
        birthday = dob_match.group(1)
    age_text = "N/A"
    age_match = re.search(r"\((\d+)\s+Year[s]?\s+Old\)", meta.subject_dob or "")
    if age_match:
        age_text = age_match.group(1)
    generated_display = meta.generated_at.strftime("%Y %b %d %H:%M")

    patient = doc.add_paragraph()
    patient.paragraph_format.space_before = Pt(4)
    patient.paragraph_format.space_after = Pt(0)
    patient_name = patient.add_run(meta.subject_full_name or meta.subject_id)
    patient_name.bold = True
    patient_name.font.size = Pt(12)

    line = doc.add_paragraph()
    line.paragraph_format.space_before = Pt(0)
    line.paragraph_format.space_after = Pt(0)
    line.add_run(f"MRN : {meta.subject_external_id}   Phone : N/A")

    line = doc.add_paragraph()
    line.paragraph_format.space_before = Pt(0)
    line.paragraph_format.space_after = Pt(0)
    line.add_run(f"Birthday : {birthday}")

    line = doc.add_paragraph()
    line.paragraph_format.space_before = Pt(0)
    line.paragraph_format.space_after = Pt(0)
    line.add_run(
        f"Visited on: N/A (Age at visit: {age_text} years) "
        f"Electronically signed by: N/A on {generated_display}"
    )

    _add_horizontal_rule(doc)

    for item in items:
        section_table = doc.add_table(rows=1, cols=2)
        section_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        section_table.autofit = False
        section_table.columns[0].width = Inches(1.95)
        section_table.columns[1].width = Inches(5.05)
        _remove_table_borders(section_table)

        title_cell = section_table.cell(0, 0)
        title_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        _set_docx_cell_margins(title_cell, top=0, bottom=0, left=0, right=90)
        title_para = title_cell.paragraphs[0]
        title_para.paragraph_format.space_after = Pt(0)
        title_para.paragraph_format.space_before = Pt(0)
        title_text = _normalize_word_text(item.title)
        title_text = title_text.replace("Patient Medical History", "Patient Medical\nHistory")
        title_run = title_para.add_run(title_text)
        title_run.font.size = Pt(13)
        title_run.bold = False

        content_cell = section_table.cell(0, 1)
        content_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        _set_docx_cell_margins(content_cell, top=0, bottom=0, left=120, right=0)
        lines = _item_word_lines(item)
        for idx, line in enumerate(lines):
            paragraph = content_cell.paragraphs[0] if idx == 0 else content_cell.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.left_indent = Pt(4)
            _append_word_line(paragraph, line)

        _set_docx_cell_border(
            title_cell,
            bottom={"val": "single", "sz": "6", "space": "0", "color": "A6A6A6"},
        )
        _set_docx_cell_border(
            content_cell,
            bottom={"val": "single", "sz": "6", "space": "0", "color": "A6A6A6"},
        )
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after = Pt(0)

    footer = section.footer.paragraphs[0]
    footer.text = f"Printed on: {generated_display}"
    footer.paragraph_format.space_after = Pt(0)
    footer2 = section.footer.add_paragraph()
    footer2.text = "Note created using Tebra"
    footer2.paragraph_format.space_before = Pt(0)
    footer2.paragraph_format.space_after = Pt(0)

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx))


def attach_file_to_pdf(pdf_path: Path, attachment_path: Path) -> None:
    reader = PdfReader(str(pdf_path))
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    if reader.metadata:
        writer.add_metadata(reader.metadata)

    with attachment_path.open("rb") as attachment_file:
        writer.add_attachment(attachment_path.name, attachment_file.read())

    temp_pdf_path = pdf_path.with_name(pdf_path.stem + ".tmp.pdf")
    with temp_pdf_path.open("wb") as out_file:
        writer.write(out_file)
    temp_pdf_path.replace(pdf_path)


def checklist_sort_key(item: Dict[str, Any]) -> tuple[float, str]:
    raw_sort_order = item.get("sortOrder")
    try:
        sort_order = float(raw_sort_order)
    except Exception:
        sort_order = 10_000.0
    display_name = str(item.get("name") or item.get("id") or "").strip().lower()
    return sort_order, display_name


def fetch_checklist_catalog(client: XcuresApiClient) -> List[Dict[str, Any]]:
    payload = client.request_json("GET", "/api/v1/patient-registry/checklist")
    candidates: List[Any]
    if isinstance(payload, list):
        candidates = payload
    elif isinstance(payload, dict):
        results = payload.get("results")
        if isinstance(results, list):
            candidates = results
        else:
            candidates = []
    else:
        raise RuntimeError(f"Unexpected checklist catalog payload type: {type(payload)}")

    checklists: List[Dict[str, Any]] = []
    for idx, item in enumerate(candidates):
        if not isinstance(item, dict):
            raise RuntimeError(f"Checklist catalog entry[{idx}] is not an object: {type(item)}")
        checklists.append(item)

    checklists.sort(key=checklist_sort_key)
    return checklists


def checklist_display_name(item: Dict[str, Any]) -> str:
    name = str(item.get("name") or "").strip()
    if name:
        return name
    item_id = str(item.get("id") or "").strip()
    return item_id or "Unnamed Checklist"


def resolve_checklist_name(checklists: Sequence[Dict[str, Any]], checklist_id: str) -> str:
    for item in checklists:
        if str(item.get("id") or "").strip() == checklist_id:
            return checklist_display_name(item)
    return checklist_id


def get_subject_header_info(client: XcuresApiClient, subject_id: str) -> tuple[str, str, str]:
    def normalize_name(name: str) -> str:
        cleaned = " ".join((name or "").split())
        if not cleaned:
            return ""
        if cleaned.isupper():
            return cleaned.title()
        return cleaned

    payload = client.request_json("GET", f"/api/v1/patient-registry/subject/{subject_id}")
    subject_obj = require_json_object(payload, context="Subject")
    first_name = normalize_name(str(subject_obj.get("firstName") or "").strip())
    last_name = normalize_name(str(subject_obj.get("lastName") or "").strip())
    full_name = " ".join(part for part in [first_name, last_name] if part).strip()
    if not full_name:
        full_name = subject_id

    external_id = "N/A"
    raw_external_ids = subject_obj.get("externalIdentifiers")
    if isinstance(raw_external_ids, list):
        for entry in raw_external_ids:
            if not isinstance(entry, dict):
                continue
            value = str(entry.get("externalIdentifier") or "").strip()
            if value:
                external_id = value
                break

    raw_dob = str(subject_obj.get("birthDate") or "").strip()
    dob = "N/A"
    if raw_dob:
        dob_date_str = raw_dob[:10]
        try:
            dob_date = date.fromisoformat(dob_date_str)
            today = datetime.now(ZoneInfo("America/New_York")).date()
            age = today.year - dob_date.year - (
                (today.month, today.day) < (dob_date.month, dob_date.day)
            )
            age_label = "Year" if age == 1 else "Years"
            dob = f"{dob_date_str} ({age} {age_label} Old)"
        except ValueError:
            dob = dob_date_str
    return full_name, external_id, dob


def prompt_subject_id(initial_subject_id: Optional[str]) -> str:
    if isinstance(initial_subject_id, str) and initial_subject_id.strip():
        return initial_subject_id.strip()

    while True:
        value = input("Enter subjectId: ").strip()
        if value:
            return value
        print("subjectId is required.")


def prompt_checklist_selection(
    checklists: Sequence[Dict[str, Any]],
    initial_checklist_id: Optional[str],
) -> tuple[str, str]:
    if isinstance(initial_checklist_id, str) and initial_checklist_id.strip():
        selected_id = initial_checklist_id.strip()
        return selected_id, resolve_checklist_name(checklists, selected_id)

    if not checklists:
        raise RuntimeError("No checklists available for interactive selection.")

    print("\nAvailable checklists:")
    for idx, item in enumerate(checklists, start=1):
        print(f"{idx} - {checklist_display_name(item)}")

    while True:
        raw = input("Select checklist number: ").strip()
        try:
            selected_idx = int(raw)
        except ValueError:
            print("Please enter a valid number.")
            continue
        if selected_idx < 1 or selected_idx > len(checklists):
            print(f"Please enter a number between 1 and {len(checklists)}.")
            continue

        selected = checklists[selected_idx - 1]
        selected_id = str(selected.get("id") or "").strip()
        if not selected_id:
            print("Selected checklist has no id. Please choose a different checklist.")
            continue
        return selected_id, checklist_display_name(selected)


def compute_unique_document_count(payload: Dict[str, Any], items: Sequence[Dict[str, Any]]) -> int:
    base_docs: set[str] = set()
    doc_ids = payload.get("documentIds")
    if isinstance(doc_ids, list):
        base_docs = {str(doc_id) for doc_id in doc_ids if str(doc_id).strip()}

    for item in items:
        item_doc_ids = item.get("documentIds")
        if isinstance(item_doc_ids, list):
            for doc_id in item_doc_ids:
                doc_str = str(doc_id).strip()
                if doc_str:
                    base_docs.add(doc_str)

        sections = item.get("documentSections")
        if isinstance(sections, list):
            for section in sections:
                if not isinstance(section, dict):
                    continue
                doc_str = str(section.get("documentId") or "").strip()
                if doc_str:
                    base_docs.add(doc_str)

    return len(base_docs)


def evaluate_checklist(
    *,
    client: XcuresApiClient,
    checklist_id: str,
    subject_id: str,
    regenerate: bool,
) -> Dict[str, Any]:
    body: Dict[str, Any] = {"subjectId": subject_id}
    if regenerate:
        body["regenerate"] = True
    endpoint = EVALUATE_ENDPOINT_TEMPLATE.format(checklist_id=checklist_id)
    payload = client.request_json("POST", endpoint, json_body=body)
    payload_obj = require_json_object(payload, context="ChecklistEvaluationResult")
    require_json_list(payload_obj.get("items"), context="ChecklistEvaluationResult.items")
    return payload_obj


def _extract_paginated_results(payload: Any) -> List[Dict[str, Any]]:
    if isinstance(payload, list):
        candidates: List[Any] = payload
    elif isinstance(payload, dict):
        candidates = []
        for key in ("results", "items", "data"):
            value = payload.get(key)
            if isinstance(value, list):
                candidates = value
                break
    else:
        candidates = []
    return [item for item in candidates if isinstance(item, dict)]


def _extract_iso_day(value: Any) -> str:
    text = str(value or "").strip()
    match = re.match(r"^(\d{4}-\d{2}-\d{2})", text)
    return match.group(1) if match else ""


def _to_sort_timestamp(value: Any) -> float:
    text = str(value or "").strip()
    if not text:
        return float("-inf")
    normalized = text.replace("Z", "+00:00")
    try:
        parsed = datetime.fromisoformat(normalized)
        if parsed.tzinfo is None:
            parsed = parsed.replace(tzinfo=ZoneInfo("America/New_York"))
        return parsed.timestamp()
    except Exception:
        day = _extract_iso_day(text)
        if day:
            try:
                return datetime.fromisoformat(f"{day}T00:00:00").timestamp()
            except Exception:
                pass
    return float("-inf")


def fetch_latest_vitals(
    *,
    client: XcuresApiClient,
    subject_id: str,
) -> tuple[List[Dict[str, Any]], str]:
    payload = client.request_json(
        "GET",
        VITALS_ENDPOINT,
        params={
            "subjectId": subject_id,
            "pageNumber": 1,
            "pageSize": 200,
            "sortField": "date",
            "sortIsDescending": True,
        },
    )
    records = _extract_paginated_results(payload)
    if not records:
        return [], ""

    records.sort(key=lambda row: _to_sort_timestamp(row.get("vitalsDate")), reverse=True)
    latest_day = ""
    for row in records:
        latest_day = _extract_iso_day(row.get("vitalsDate"))
        if latest_day:
            break

    if latest_day:
        latest_records = [
            row for row in records if _extract_iso_day(row.get("vitalsDate")) == latest_day
        ]
    else:
        latest_records = [records[0]]

    return latest_records, latest_day


def fetch_high_labs_last_12_months(
    *,
    client: XcuresApiClient,
    subject_id: str,
    now: datetime,
) -> tuple[List[Dict[str, Any]], str, str]:
    window_end = now.date().isoformat()
    window_start = (now.date() - timedelta(days=365)).isoformat()

    records = client.list_paginated(
        LABS_ENDPOINT,
        params={
            "subjectId": subject_id,
            "dateStart": window_start,
            "dateEnd": window_end,
            "sortField": "date",
            "sortIsDescending": True,
        },
        page_size=200,
        max_pages=50,
    )
    high_records = [
        row
        for row in records
        if str(row.get("labResultInterpretation") or "").strip().lower() == "high"
    ]
    high_records.sort(key=lambda row: _to_sort_timestamp(row.get("labDate")), reverse=True)
    return high_records, window_start, window_end


def _build_supplemental_item(
    *,
    title: str,
    sort_order: float,
    lines: List[str],
    source_records: List[Dict[str, Any]],
) -> ItemViewModel:
    doc_ids: set[str] = set()
    for row in source_records:
        raw_doc_ids = row.get("documentIds")
        if isinstance(raw_doc_ids, list):
            for raw_doc_id in raw_doc_ids:
                value = str(raw_doc_id or "").strip()
                if value:
                    doc_ids.add(value)

    value = "\n".join(lines).strip() or "No data available."
    return ItemViewModel(
        title=title,
        sort_order=sort_order,
        meets_criteria=True,
        details=[DetailRow(label="Items", value=value)],
        result_raw=None,
        evidence=value,
        source_lines=[],
        document_count=len(doc_ids),
        result_is_structured=True,
    )


def _format_vital_line(row: Dict[str, Any]) -> str:
    vital_name = str(row.get("vitals") or "Vital").strip()
    result_text = str(row.get("vitalsResult") or "").strip()
    value_text = str(row.get("vitalsResultValue") or "").strip()
    unit_text = str(row.get("vitalsResultUnit") or "").strip()
    measurement = result_text or " ".join(part for part in [value_text, unit_text] if part).strip()
    if not measurement:
        measurement = "N/A"
    return f"{vital_name}: {measurement}"


def _format_lab_line(row: Dict[str, Any]) -> str:
    lab_name = _lab_display_name(row)
    if not lab_name:
        lab_name = "Unknown Lab"

    result_value = str(row.get("labResultValue") or row.get("labResult") or "").strip()
    if not result_value:
        result_value = "N/A"

    interpretation = str(row.get("labResultInterpretation") or "").strip().upper()
    if interpretation == "HIGH":
        marker = " (H)"
    elif interpretation == "LOW":
        marker = " (L)"
    else:
        marker = ""

    return f"{lab_name}: {result_value}{marker}"


def _lab_display_name(row: Dict[str, Any]) -> str:
    raw_codes = row.get("codes")
    if isinstance(raw_codes, list):
        standard_display = ""
        fallback_display = ""
        for code_obj in raw_codes:
            if not isinstance(code_obj, dict):
                continue
            display = str(code_obj.get("display") or "").strip()
            if not display:
                continue
            if not fallback_display:
                fallback_display = display
            if str(code_obj.get("type") or "").strip().lower() == "standard":
                standard_display = display
                break
        if standard_display:
            return standard_display
        if fallback_display:
            return fallback_display
    return ""


def _lab_dedupe_key(row: Dict[str, Any]) -> str:
    primary = _lab_display_name(row).strip().lower()
    if primary:
        return primary

    raw_codes = row.get("codes")
    if isinstance(raw_codes, list):
        codes = [str(code).strip().lower() for code in raw_codes if str(code).strip()]
        if codes:
            return "|".join(sorted(codes))

    fallback = str(row.get("id") or "").strip().lower()
    return fallback or "unknown-lab"


def _supplemental_sort_orders(items: Sequence[ItemViewModel]) -> tuple[float, float]:
    social_sort_order = next(
        (item.sort_order for item in items if item.title.strip().lower() == "social history"),
        None,
    )
    if social_sort_order is not None:
        return social_sort_order - 0.20, social_sort_order - 0.10

    max_sort_order = max((item.sort_order for item in items), default=0.0)
    return max_sort_order + 0.10, max_sort_order + 0.20


def _upsert_section_item(items: List[ItemViewModel], section_item: ItemViewModel) -> None:
    section_title = section_item.title.strip().lower()
    for idx, item in enumerate(items):
        if item.title.strip().lower() == section_title:
            items[idx] = section_item
            return
    items.append(section_item)


def _latest_high_labs_per_lab(records: Sequence[Dict[str, Any]]) -> List[Dict[str, Any]]:
    latest_by_key: Dict[str, Dict[str, Any]] = {}
    for row in records:
        key = _lab_dedupe_key(row)
        existing = latest_by_key.get(key)
        if existing is None:
            latest_by_key[key] = row
            continue
        if _to_sort_timestamp(row.get("labDate")) > _to_sort_timestamp(existing.get("labDate")):
            latest_by_key[key] = row
    deduped = list(latest_by_key.values())
    deduped.sort(key=lambda row: _to_sort_timestamp(row.get("labDate")), reverse=True)
    return deduped


def build_output_paths(
    *,
    output_dir: Path,
    subject_id: str,
    checklist_id: str,
    generated_at: datetime,
) -> tuple[Path, Path, Path]:
    timestamp = generated_at.strftime("%Y%m%d_%H%M%S")
    base_name = (
        f"checklist_eval_{compact_id(subject_id)}_{compact_id(checklist_id)}_{timestamp}"
    )
    return (
        output_dir / f"{base_name}.pdf",
        output_dir / f"{base_name}.json",
        output_dir / f"{base_name}.docx",
    )


def main() -> int:
    load_env_file(REPO_ROOT / ".env")
    args = parse_args()

    output_dir = DEFAULT_OUTPUT_DIR.expanduser().resolve()
    output_dir.mkdir(parents=True, exist_ok=True)

    bearer = args.bearer.strip() if isinstance(args.bearer, str) and args.bearer.strip() else None
    if not bearer:
        bearer = get_xcures_bearer_token(timeout_seconds=args.timeout)

    config = build_render_config()
    now = datetime.now(ZoneInfo("America/New_York"))
    latest_vital_records: List[Dict[str, Any]] = []
    latest_vital_day = ""
    vitals_error = ""

    high_lab_records: List[Dict[str, Any]] = []
    labs_window_start = (now.date() - timedelta(days=365)).isoformat()
    labs_window_end = now.date().isoformat()
    labs_error = ""

    with requests.Session() as session:
        client = XcuresApiClient(
            session=session,
            base_url=str(args.base_url).rstrip("/"),
            project_id=args.project_id,
            bearer_token=bearer,
            timeout_seconds=args.timeout,
        )

        checklist_catalog: List[Dict[str, Any]]
        try:
            checklist_catalog = fetch_checklist_catalog(client)
        except Exception as exc:
            if not args.checklist_id:
                raise RuntimeError(
                    "Failed to load checklist catalog for interactive selection. "
                    "Provide --checklist-id to bypass selection."
                ) from exc
            print(
                f"Warning: failed to retrieve checklist catalog; using checklist ID as title ({exc})",
                file=sys.stderr,
            )
            checklist_catalog = []

        subject_id = prompt_subject_id(args.subject_id)
        subject_full_name = subject_id
        subject_external_id = "N/A"
        subject_dob = "N/A"
        try:
            subject_full_name, subject_external_id, subject_dob = get_subject_header_info(
                client, subject_id
            )
        except Exception as exc:
            print(
                "Warning: failed to retrieve subject header info; "
                f"using fallbacks ({exc})",
                file=sys.stderr,
            )

        checklist_id, checklist_name = prompt_checklist_selection(
            checklist_catalog,
            args.checklist_id,
        )

        evaluation = evaluate_checklist(
            client=client,
            checklist_id=checklist_id,
            subject_id=subject_id,
            regenerate=bool(args.regenerate),
        )
        try:
            latest_vital_records, latest_vital_day = fetch_latest_vitals(
                client=client,
                subject_id=subject_id,
            )
        except Exception as exc:
            vitals_error = str(exc)
            print(
                "Warning: failed to fetch latest vitals; section will include error "
                f"details ({exc})",
                file=sys.stderr,
            )
        try:
            high_lab_records, labs_window_start, labs_window_end = fetch_high_labs_last_12_months(
                client=client,
                subject_id=subject_id,
                now=now,
            )
        except Exception as exc:
            labs_error = str(exc)
            print(
                "Warning: failed to fetch high labs; section will include error "
                f"details ({exc})",
                file=sys.stderr,
            )

    raw_items = require_json_list(evaluation.get("items"), context="ChecklistEvaluationResult.items")
    typed_items: List[Dict[str, Any]] = []
    for idx, item in enumerate(raw_items):
        if not isinstance(item, dict):
            raise RuntimeError(
                f"ChecklistEvaluationResult.items[{idx}] must be an object; got {type(item)}"
            )
        typed_items.append(item)
    normalized_items = [normalize_item(item) for item in typed_items]
    normalized_items = [
        ItemViewModel(
            title=clean_section_title(item.title, checklist_name),
            sort_order=item.sort_order,
            meets_criteria=item.meets_criteria,
            details=item.details,
            result_raw=item.result_raw,
            evidence=item.evidence,
            source_lines=item.source_lines,
            document_count=item.document_count,
            result_is_structured=item.result_is_structured,
        )
        for item in normalized_items
    ]

    vitals_sort_order, labs_sort_order = _supplemental_sort_orders(normalized_items)
    if vitals_error:
        vitals_lines = [f"Unable to retrieve vitals: {vitals_error}"]
    elif not latest_vital_records:
        vitals_lines = ["No vitals found."]
    else:
        vitals_lines = [f"Recorded: {latest_vital_day}"] if latest_vital_day else []
        seen_vitals: set[str] = set()
        for record in latest_vital_records:
            line = _format_vital_line(record)
            if line in seen_vitals:
                continue
            seen_vitals.add(line)
            vitals_lines.append(line)
    vitals_item = _build_supplemental_item(
        title="Vitals",
        sort_order=vitals_sort_order,
        lines=vitals_lines,
        source_records=latest_vital_records,
    )

    deduped_high_lab_records = _latest_high_labs_per_lab(high_lab_records)
    if labs_error:
        labs_lines = [f"Unable to retrieve labs: {labs_error}"]
    elif not deduped_high_lab_records:
        labs_lines = [f"No high lab results from {labs_window_start} to {labs_window_end}."]
    else:
        labs_lines = [_format_lab_line(record) for record in deduped_high_lab_records]
    labs_item = _build_supplemental_item(
        title="Labs",
        sort_order=labs_sort_order,
        lines=labs_lines,
        source_records=deduped_high_lab_records,
    )

    _upsert_section_item(normalized_items, vitals_item)
    _upsert_section_item(normalized_items, labs_item)
    normalized_items.sort(key=lambda item: (item.sort_order, item.title.lower()))

    total_items = len(normalized_items)
    passed_items = sum(1 for item in normalized_items if item.meets_criteria)
    failed_items = total_items - passed_items
    unique_document_count = compute_unique_document_count(evaluation, typed_items)
    eligibility_satisfied = bool(evaluation.get("eligibilitySatisfied"))

    meta = ReportMeta(
        subject_id=subject_id,
        subject_full_name=subject_full_name,
        subject_external_id=subject_external_id,
        subject_dob=subject_dob,
        checklist_id=checklist_id,
        checklist_name=checklist_name,
        eligibility_satisfied=eligibility_satisfied,
        generated_at=now,
        total_items=total_items,
        passed_items=passed_items,
        failed_items=failed_items,
        unique_document_count=unique_document_count,
        endpoint_path="POST /api/v1/patient-registry/checklist/{checklistId}/evaluate",
    )

    pdf_path, json_path, docx_path = build_output_paths(
        output_dir=output_dir,
        subject_id=subject_id,
        checklist_id=checklist_id,
        generated_at=now,
    )

    write_pdf_report(meta=meta, items=normalized_items, output_pdf=pdf_path, config=config)
    if args.save_docx:
        write_docx_report(meta=meta, items=normalized_items, output_docx=docx_path)

    if args.save_json:
        with json_path.open("w", encoding="utf-8") as out_file:
            json.dump(evaluation, out_file, indent=2, ensure_ascii=False, sort_keys=True)
        attach_file_to_pdf(pdf_path, json_path)

    print(f"Wrote PDF: {pdf_path}")
    if args.save_docx:
        print(f"Wrote DOCX: {docx_path}")
    if args.save_json:
        print(f"Wrote JSON: {json_path}")
        print("Attached JSON to PDF as embedded file.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
